from __future__ import annotations

import dataclasses
from decimal import Decimal
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from backend.schemas.extraction import ExtractionResult
from backend.services.template_service import filter_extraction_by_template

# ── Colour palette ─────────────────────────────────────────────────────────────
_BLUE = RGBColor(0x2D, 0x5B, 0x8A)
_RED = RGBColor(0xCC, 0x00, 0x00)
_GREY = RGBColor(0x66, 0x66, 0x66)

# Anchor fields considered "fiscal summary"
_FISCAL_FIELDS = [
    "base_imponible",
    "iva_rate",
    "iva_amount",
    "irpf_rate",
    "irpf_amount",
    "total_amount",
    "currency",
]

_ISSUER_FIELDS = ["issuer_name", "issuer_cif"]
_RECIPIENT_FIELDS = ["recipient_name", "recipient_cif"]


# ── Helpers ────────────────────────────────────────────────────────────────────

def _safe(value: Any) -> str:
    """Convert any value to a display string; None → empty string."""
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return f"{value:,.2f}"
    return str(value)


def _add_table_row(table: Any, label: str, value: str, monetary: bool = False) -> None:
    """Append a row to a 2-column table: bold label | right-aligned value."""
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_para = label_cell.paragraphs[0]
    run = label_para.add_run(label)
    run.bold = True

    value_para = value_cell.paragraphs[0]
    value_para.add_run(value)
    if monetary:
        value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _build_fiscal_table(doc: Document, anchor: Any) -> None:
    """Write the 'Resumen Fiscal' 2-column table."""
    anchor_dict = dataclasses.asdict(anchor)
    fiscal_rows: list[tuple[str, str, bool]] = []
    for key in _FISCAL_FIELDS:
        value = anchor_dict.get(key)
        if value is None:
            continue
        label = key.replace("_", " ").title()
        display = _safe(value)
        is_monetary = isinstance(value, Decimal)
        fiscal_rows.append((label, display, is_monetary))

    if not fiscal_rows:
        doc.add_paragraph("No se encontraron valores fiscales.", style="Normal")
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Campo"
    hdr_cells[1].text = "Valor"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for label, display, is_monetary in fiscal_rows:
        _add_table_row(table, label, display, monetary=is_monetary)


def _build_parties_section(doc: Document, anchor: Any) -> None:
    """Write issuer and recipient blocks."""
    doc.add_heading("Partes", level=2)
    anchor_dict = dataclasses.asdict(anchor)

    issuer_name = anchor_dict.get("issuer_name") or ""
    issuer_cif = anchor_dict.get("issuer_cif") or ""
    recipient_name = anchor_dict.get("recipient_name") or ""
    recipient_cif = anchor_dict.get("recipient_cif") or ""

    if issuer_name or issuer_cif:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run("Emisor: ")
        run.bold = True
        p.add_run(f"{issuer_name}  |  CIF: {issuer_cif}")

    if recipient_name or recipient_cif:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run("Receptor: ")
        run.bold = True
        p.add_run(f"{recipient_name}  |  CIF: {recipient_cif}")


def _build_additional_details(doc: Document, result: ExtractionResult) -> None:
    """Write discovered fields that are not line_items or raw_text."""
    discovered = result.discovered or {}
    extra: dict[str, Any] = {
        k: v
        for k, v in discovered.items()
        if k not in ("line_items", "raw_text")
    }
    if not extra:
        return

    doc.add_heading("Detalles Adicionales", level=2)
    for key, value in extra.items():
        if value is None:
            continue
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(f"{key.replace('_', ' ').title()}: ")
        run.bold = True
        p.add_run(_safe(value))


def _build_issues_section(doc: Document, result: ExtractionResult) -> None:
    """Write the 'Incidencias' section if issues are present."""
    issues = result.issues or []
    if not issues:
        return

    doc.add_heading("Incidencias", level=2)
    for issue in issues:
        severity_label = f"[{issue.severity.upper()}]"
        field_part = f" ({issue.field})" if issue.field else ""
        text = f"{severity_label}{field_part}: {issue.message}"
        p = doc.add_paragraph(text, style="List Bullet")
        # Colour prefix by severity
        if p.runs:
            run = p.runs[0]
            if issue.severity == "error":
                run.font.color.rgb = _RED
            elif issue.severity == "warning":
                run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)


# ── Public API ─────────────────────────────────────────────────────────────────

def to_docx(
    result: ExtractionResult,
    filename: str,
    template_fields: list[dict] | None = None,
) -> bytes:
    """
    Export ExtractionResult to Word bytes.

    Structure:
    - Title: "Factura {invoice_number} — {issuer_name}"
    - Metadata line: "Fecha: {date}  |  Emisor CIF: {cif}  |  Receptor: {name}"
    - Heading 2: "Resumen Fiscal"
    - 2-column table: fiscal fields
    - Heading 2: "Partes" with issuer / recipient blocks
    - Heading 2: "Detalles Adicionales" (only if discovered has non-standard fields)
    - Heading 2: "Incidencias" (only if result.issues is non-empty)
    - Footer: "Extraído con {model} el {timestamp}"
    """
    doc = Document()
    anchor = result.anchor

    # ── Title ────────────────────────────────────────────────────────────────
    invoice_number = getattr(anchor, "invoice_number", None) or ""
    issuer_name = getattr(anchor, "issuer_name", None) or ""
    title_text = f"Factura {invoice_number} — {issuer_name}".strip(" —")
    doc.add_paragraph(title_text, style="Title")

    # ── Metadata line ────────────────────────────────────────────────────────
    issue_date = getattr(anchor, "issue_date", None) or ""
    issuer_cif = getattr(anchor, "issuer_cif", None) or ""
    recipient_name = getattr(anchor, "recipient_name", None) or ""
    p_meta = doc.add_paragraph(style="Normal")
    run_meta = p_meta.add_run(f"Fecha: {issue_date}  |  Emisor CIF: {issuer_cif}  |  Receptor: {recipient_name}")
    run_meta.font.color.rgb = _GREY

    # ── Resumen Fiscal ───────────────────────────────────────────────────────
    doc.add_heading("Resumen Fiscal", level=2)
    if template_fields:
        filtered = filter_extraction_by_template(result, template_fields)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Campo"
        hdr_cells[1].text = "Valor"
        for cell in hdr_cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for display_name, value in filtered.items():
            is_monetary = isinstance(value, Decimal)
            _add_table_row(table, display_name, _safe(value), monetary=is_monetary)
    else:
        _build_fiscal_table(doc, anchor)

    # ── Partes ───────────────────────────────────────────────────────────────
    _build_parties_section(doc, anchor)

    # ── Detalles Adicionales ─────────────────────────────────────────────────
    _build_additional_details(doc, result)

    # ── Incidencias ──────────────────────────────────────────────────────────
    _build_issues_section(doc, result)

    # ── Footer ───────────────────────────────────────────────────────────────
    model = result.llm_model or "unknown"
    timestamp = result.extraction_timestamp or ""
    p_footer = doc.add_paragraph(style="Normal")
    run_footer = p_footer.add_run(f"Extraído con {model} el {timestamp}")
    run_footer.font.color.rgb = _GREY
    run_footer.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
