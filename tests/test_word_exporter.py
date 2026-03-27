from __future__ import annotations

from decimal import Decimal
from io import BytesIO

import pytest
from docx import Document

from backend.schemas.extraction import AnchorFields, ExtractionIssue, ExtractionResult
from backend.services.word_exporter import to_docx


def _make_result(issues=None) -> ExtractionResult:
    anchor = AnchorFields(
        issuer_name="INMOBILIARIA BARUAL S.L.",
        issuer_cif="B50042332",
        recipient_name="GESTINAD ARAGON S.L.",
        recipient_cif="B99334708",
        invoice_number="GAT143/26",
        issue_date="2026-03-01",
        base_imponible=Decimal("102.57"),
        iva_rate=Decimal("21"),
        iva_amount=Decimal("21.54"),
        irpf_rate=None,
        irpf_amount=None,
        total_amount=Decimal("124.11"),
        currency="EUR",
    )
    return ExtractionResult(
        anchor=anchor,
        discovered={
            "line_items": [{"concept": "RENTA LEGAL", "total_amount": Decimal("102.57")}],
            "raw_text": "Sample raw text.",
        },
        issues=issues
        or [
            ExtractionIssue(
                field=None,
                message="IVA looks correct",
                severity="observation",
                source="llm",
            )
        ],
        requires_review=False,
        llm_model="qwen3.5:9b",
        extraction_timestamp="2026-03-01T10:00:00Z",
    )


def _all_text(doc: Document) -> str:
    """Collect all paragraph text in the document."""
    return "\n".join(p.text for p in doc.paragraphs)


# ── Test 1 ────────────────────────────────────────────────────────────────────
def test_to_docx_returns_bytes():
    result = _make_result()
    output = to_docx(result, "invoice.pdf")
    assert isinstance(output, bytes)
    assert len(output) > 0


# ── Test 2 ────────────────────────────────────────────────────────────────────
def test_to_docx_title_contains_invoice_number():
    """The first Title-style paragraph should contain the invoice number."""
    result = _make_result()
    output = to_docx(result, "invoice.pdf")
    doc = Document(BytesIO(output))
    full_text = _all_text(doc)
    assert "GAT143/26" in full_text


# ── Test 3 ────────────────────────────────────────────────────────────────────
def test_to_docx_issues_section_present_when_issues():
    """When there are issues, the document must contain the 'Incidencias' heading."""
    issues = [
        ExtractionIssue(
            field="iva_amount",
            message="IVA arithmetic mismatch: expected 21.54, got 21.50",
            severity="error",
            source="validator",
        ),
        ExtractionIssue(
            field=None,
            message="Recipient CIF missing",
            severity="warning",
            source="validator",
        ),
    ]
    result = _make_result(issues=issues)
    output = to_docx(result, "invoice.pdf")
    doc = Document(BytesIO(output))
    full_text = _all_text(doc)
    assert "Incidencias" in full_text
